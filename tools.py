"""
General-purpose PDF utilities beyond signing: merge, extract pages,
compress, fill forms, convert to Word, and password protect/unlock.
"""

import io
import os
import shutil
import subprocess
import tempfile

import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import BooleanObject, NameObject


# ---------------------------------------------------------------------------
# Merge
# ---------------------------------------------------------------------------

def merge_pdfs(pdf_byte_list):
    """Merge PDFs (as raw bytes) in the given order into one PDF."""
    writer = PdfWriter()
    for pdf_bytes in pdf_byte_list:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Extract pages
# ---------------------------------------------------------------------------

def parse_page_spec(spec: str, page_count: int):
    """Parse a 1-based page spec like '1,3,5-8' into a deduped, ordered list
    of 0-based page indices, validated against page_count."""
    indices = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            a, b = int(a.strip()), int(b.strip())
            if a > b:
                a, b = b, a
            indices.extend(range(a, b + 1))
        else:
            indices.append(int(part))

    if not indices:
        raise ValueError("No pages specified")

    seen = set()
    result = []
    for one_based in indices:
        zero_based = one_based - 1
        if zero_based < 0 or zero_based >= page_count:
            raise ValueError(f"Page {one_based} is out of range (this PDF has {page_count} pages)")
        if zero_based not in seen:
            seen.add(zero_based)
            result.append(zero_based)
    return result


def extract_pages(pdf_bytes: bytes, page_spec: str) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    indices = parse_page_spec(page_spec, len(reader.pages))

    writer = PdfWriter()
    for i in indices:
        writer.add_page(reader.pages[i])

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Compress
# ---------------------------------------------------------------------------

_GS_QUALITY_PRESETS = {"screen", "ebook", "printer", "prepress"}


def compress_pdf(pdf_bytes: bytes, quality: str = "ebook") -> bytes:
    """
    Shrink a PDF's file size. Prefers Ghostscript (downsamples embedded
    images - much bigger savings) and falls back to qpdf's stream
    compression/optimization if Ghostscript isn't installed.
    """
    if quality not in _GS_QUALITY_PRESETS:
        quality = "ebook"

    if shutil.which("gs"):
        return _compress_with_ghostscript(pdf_bytes, quality)
    if shutil.which("qpdf"):
        return _compress_with_qpdf(pdf_bytes)
    raise RuntimeError("Neither ghostscript nor qpdf is available on this system for compression")


def _compress_with_ghostscript(pdf_bytes: bytes, quality: str) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f_in:
        f_in.write(pdf_bytes)
        in_path = f_in.name
    out_path = in_path + "_out.pdf"

    try:
        cmd = [
            "gs",
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            f"-dPDFSETTINGS=/{quality}",
            "-dNOPAUSE",
            "-dBATCH",
            "-dQUIET",
            f"-sOutputFile={out_path}",
            in_path,
        ]
        result = subprocess.run(cmd, capture_output=True, timeout=120)
        if result.returncode != 0 or not os.path.exists(out_path):
            raise RuntimeError(f"Ghostscript failed: {result.stderr.decode(errors='ignore')[:500]}")
        with open(out_path, "rb") as f:
            return f.read()
    finally:
        for p in (in_path, out_path):
            if os.path.exists(p):
                os.remove(p)


def _compress_with_qpdf(pdf_bytes: bytes) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f_in:
        f_in.write(pdf_bytes)
        in_path = f_in.name
    out_path = in_path + "_out.pdf"

    try:
        cmd = ["qpdf", "--compress-streams=y", "--object-streams=generate", "--optimize-images", in_path, out_path]
        result = subprocess.run(cmd, capture_output=True, timeout=120)
        # qpdf returns 0 for full success, 3 for "warnings but usable output"
        if result.returncode not in (0, 3) or not os.path.exists(out_path):
            raise RuntimeError(f"qpdf failed: {result.stderr.decode(errors='ignore')[:500]}")
        with open(out_path, "rb") as f:
            return f.read()
    finally:
        for p in (in_path, out_path):
            if os.path.exists(p):
                os.remove(p)


# ---------------------------------------------------------------------------
# Fill forms (AcroForm fields only - PDFs without real form fields aren't
# supported here; the signature-placement drag/drop flow is the fallback
# for stamping arbitrary text onto a non-fillable PDF)
# ---------------------------------------------------------------------------

_FIELD_TYPE_MAP = {"/Tx": "text", "/Btn": "checkbox", "/Ch": "choice"}


def get_form_fields(pdf_bytes: bytes):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    fields = reader.get_fields()
    if not fields:
        return []

    out = []
    for name, f in fields.items():
        raw_type = f.get("/FT")
        field_type = _FIELD_TYPE_MAP.get(str(raw_type), "text")
        entry = {
            "name": name,
            "type": field_type,
            "value": str(f.get("/V")) if f.get("/V") is not None else "",
        }
        if field_type == "choice":
            opts = f.get("/Opt")
            if opts:
                entry["options"] = [str(o) for o in opts]
        out.append(entry)
    return out


def fill_form(pdf_bytes: bytes, values: dict) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    writer.append(reader)

    # Ask viewers to regenerate field appearances so typed/selected values
    # actually render visibly rather than showing blank.
    if "/AcroForm" in writer._root_object:
        writer._root_object["/AcroForm"][NameObject("/NeedAppearances")] = BooleanObject(True)

    for page in writer.pages:
        writer.update_page_form_field_values(page, values)

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Convert to Word (text + table extraction - not a pixel-perfect layout
# clone, but a genuinely editable .docx)
# ---------------------------------------------------------------------------

def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    doc = Document()

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            for line in text.split("\n"):
                doc.add_paragraph(line)

            for table in page.extract_tables():
                if not table:
                    continue
                rows = len(table)
                cols = max(len(r) for r in table)
                docx_table = doc.add_table(rows=rows, cols=cols)
                docx_table.style = "Table Grid"
                for r_idx, row in enumerate(table):
                    for c_idx in range(cols):
                        cell_value = row[c_idx] if c_idx < len(row) and row[c_idx] else ""
                        docx_table.rows[r_idx].cells[c_idx].text = str(cell_value)

            if i < page_count - 1:
                doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Password protect / unlock
# ---------------------------------------------------------------------------

def protect_pdf(pdf_bytes: bytes, user_password: str, owner_password: str = "") -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password=user_password, owner_password=owner_password or user_password)

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def unlock_pdf(pdf_bytes: bytes, password: str) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        # pypdf's decrypt() does NOT raise on a wrong password - it returns
        # PasswordType.NOT_DECRYPTED (falsy) instead, so we have to check
        # the return value explicitly.
        result = reader.decrypt(password)
        if not result:
            raise ValueError("Incorrect password")

    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Organize (reorder / rotate / delete pages)
# ---------------------------------------------------------------------------

def get_page_thumbnails(pdf_bytes: bytes, dpi: int = 90):
    """Render each page as a small preview PNG for a reorder/rotate/delete UI."""
    import base64

    pdf = pdfium.PdfDocument(io.BytesIO(pdf_bytes))
    thumbnails = []
    try:
        scale = dpi / 72.0
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=scale)
            pil_image = bitmap.to_pil().convert("RGB")
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG")
            thumbnails.append(
                {"index": i, "image_base64": base64.b64encode(buf.getvalue()).decode("ascii")}
            )
    finally:
        pdf.close()
    return thumbnails


def organize_pdf(pdf_bytes: bytes, page_order, rotations=None) -> bytes:
    """
    Rebuild a PDF from a subset/reordering of its original pages, with
    optional per-original-index rotation (in degrees, added to whatever
    rotation the page already had).

    page_order: list of original 0-based page indices, in the desired new
        order. Omitted indices are dropped (deleted). May repeat an index
        to duplicate a page.
    rotations: dict of {original_index: degrees}, degrees a multiple of 90.
    """
    rotations = rotations or {}
    reader = PdfReader(io.BytesIO(pdf_bytes))
    page_count = len(reader.pages)

    for idx in page_order:
        if idx < 0 or idx >= page_count:
            raise ValueError(f"Page index {idx} is out of range (this PDF has {page_count} pages)")
    if not page_order:
        raise ValueError("No pages selected - the result would be an empty PDF")

    writer = PdfWriter()
    for idx in page_order:
        page = reader.pages[idx]
        degrees = rotations.get(str(idx), rotations.get(idx, 0))
        if degrees:
            page.rotate(int(degrees))
        writer.add_page(page)

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()
